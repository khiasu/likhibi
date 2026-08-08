"""
Single Authoritative Script to Generate the Mini Project Presentation DOCX.

Strictly follows the University ML/AI Project Review Guidelines PDF:
  1. Aim
  2. Objectives
  3. Motivation
  4. Literature Survey & Gap Analysis
  5. Problem Statement
  6. Proposed Methodology (ML Pipeline: Data Collection, Pre-processing,
     Feature Engineering, Dataset Splitting, Model Selection, Model Training,
     Model Evaluation, Prediction/Deployment, Scope Split, Step-by-Step Workflow)
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, title, text, bg_hex="F0F4F8", border_hex="1A365D"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"🗣️ {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.italic = True
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_mini_project_docx(output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = doc.styles
    normal_font = styles['Normal'].font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("MINI PROJECT REVIEW – I (PROPOSAL DEFENSE)")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run("ML/AI Guideline-Aligned Presentation Slide Deck Script & Defense Speaker Notes\nLikhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Meta Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Review Type:", "Mini Project Proposal Defense (Project Review – I)"),
        ("ML Pipeline Focus:", "Corpus Curation + 20k Lexicon + N-Gram Model + Trie Index + Android IME Demo APK"),
        ("Project Coordinators:", "Mr. Nzanthung Odyuo | Mr. Nokshangthemba")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width, cell_v.width = Inches(2.0), Inches(4.5)
        set_cell_background(cell_k, "F7FAFC")
        set_cell_background(cell_v, "FFFFFF")
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(2)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(9.5)
        rk.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(2)
        rv = pv.add_run(v)
        rv.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Scope Split Visual Box
    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_table.autofit = False
    mini_cell = scope_table.cell(0, 0)
    major_cell = scope_table.cell(0, 1)
    mini_cell.width = Inches(3.2)
    major_cell.width = Inches(3.3)

    set_cell_background(mini_cell, "EBF8FF")
    set_cell_margins(mini_cell, top=120, bottom=120, left=150, right=100)
    tcPr_m = mini_cell._element.get_or_add_tcPr()
    borders_m = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2196F3"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr_m.append(borders_m)
    p_m = mini_cell.paragraphs[0]
    r_m1 = p_m.add_run("✅ MINI PROJECT DELIVERABLES (THIS REVIEW)\n")
    r_m1.bold = True; r_m1.font.size = Pt(9.5); r_m1.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    r_m2 = p_m.add_run("• Monolingual & Parallel Corpus\n• 20,000+ Entry Lexical Database\n• N-Gram Language Model\n• Character-Level Trie Index\n• Initial Android IME Keyboard APK")
    r_m2.font.size = Pt(9.5)

    set_cell_background(major_cell, "FFF3E0")
    set_cell_margins(major_cell, top=120, bottom=120, left=150, right=100)
    tcPr_ma = major_cell._element.get_or_add_tcPr()
    borders_ma = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="F57C00"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr_ma.append(borders_ma)
    p_ma = major_cell.paragraphs[0]
    r_ma1 = p_ma.add_run("🔜 MAJOR PROJECT CONTINUATION\n")
    r_ma1.bold = True; r_ma1.font.size = Pt(9.5); r_ma1.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    r_ma2 = p_ma.add_run("• Neural Machine Translation (NMT)\n• Nagamese ↔ English seq2seq model\n• Final polished APK release\n• Complete research documentation\n• BLEU score translation evaluation")
    r_ma2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 10 SLIDES STRUCTURE FOLLOWING ML GUIDELINES PDF
    slides = [
        {
            "num": "SLIDE 1",
            "title": "Title Slide",
            "bullets": [
                "Project Title: Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole",
                "Review Level: Mini Project Proposal Defense (Project Review – I)",
                "Domain: Natural Language Processing / Machine Learning",
                "Target Language: Nagamese Creole (Lingua Franca of Nagaland)",
                "Presenter: [Your Name] | Roll No / USN: [Your Roll Number] | Dept of CSE",
                "Project Coordinators: Mr. Nzanthung Odyuo & Mr. Nokshangthemba"
            ],
            "notes": "Respected Project Coordinators Mr. Nzanthung Odyuo sir, Mr. Nokshangthemba sir, and evaluation committee members. Good morning. I am [Your Name], presenting my mini project proposal titled 'Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole'. Nagamese is the primary spoken lingua franca across Nagaland with over 30 million daily users, yet it remains an extremely Low-Resource Language in computer science. This mini project constructs foundational NLP datasets and language models, deploying a functional initial Android keyboard as our demonstration platform."
        },
        {
            "num": "SLIDE 2",
            "title": "Table of Contents",
            "bullets": [
                "1. Aim & Specific Mini Project Objectives",
                "2. Motivation & Significance",
                "3. Literature Survey & Gap Analysis",
                "4. Problem Statement",
                "5. Proposed Methodology – ML Architecture & Mini/Major Scope Split",
                "6. Proposed Methodology – Stage 1: Data Collection & Stage 2: Data Pre-processing",
                "7. Proposed Methodology – Stage 3: Feature Engineering, Stage 4: Dataset Splitting & Stage 5: Model Selection",
                "8. Proposed Methodology – Stage 6: Model Training, Stage 7: Model Evaluation & Stage 8: Prediction / Deployment",
                "9. Step-by-Step ML Workflow",
                "10. Expected Outcomes, Evaluation Metrics & Key References"
            ],
            "notes": "Here is the 10-slide table of contents structured strictly according to the University ML/AI Project Review guidelines. I will cover our aim and objectives, establish our motivation, survey literature gaps, define the problem statement, and detail our end-to-end Machine Learning pipeline across data collection, preprocessing, feature extraction, model selection, training, evaluation, and mobile deployment."
        },
        {
            "num": "SLIDE 3",
            "title": "Aim & Objectives (ML Guideline Items 1 & 2)",
            "bullets": [
                "1. AIM OF THE PROJECT (Item 1):",
                "   'To design, curate, and implement a foundational NLP resource infrastructure for Nagamese Creole, train an offline contextual word-prediction language model, and deploy a functional initial Android Input Method Editor (IME) keyboard application.'",
                "2. SPECIFIC MINI PROJECT OBJECTIVES (Item 2):",
                "   • Corpus Building: Construct a 6,965-line monolingual corpus & 6,965-pair English-Nagamese parallel corpus.",
                "   • Lexical Database: Compile a verified 20,000+ entry Nagamese digital dictionary (`nagamese_lexicon.json`).",
                "   • Language Modeling: Train statistical N-gram models (Unigram, Bigram, Trigram) with add-k smoothing.",
                "   • Trie Indexing: Construct a character-level Trie prefix tree index for sub-millisecond prefix autocompletion.",
                "   • Android IME Deployment: Integrate models into an offline Android keyboard APK operating under 5 MB RAM."
            ],
            "notes": "Slide 3 defines our Aim and Objectives following ML guideline items 1 and 2. Our Aim is to build the foundational computational resources for Nagamese and deploy an offline predictive text engine. For the Mini Project, we have 5 concrete objectives: corpus creation, a 20,000-entry validated lexicon, statistical N-gram language modeling, Trie index construction, and an initial functional Android keyboard application."
        },
        {
            "num": "SLIDE 4",
            "title": "Motivation & Significance (ML Guideline Item 3)",
            "bullets": [
                "Sociolinguistic Reality:",
                "  • Nagamese is spoken by over 30 million people across Nagaland and neighboring regions as an inter-tribal lingua franca.",
                "The Low-Resource Language Crisis:",
                "  • Despite massive daily messaging, major mobile operating systems (Android/iOS) offer ZERO native predictive text support for Nagamese in Roman script.",
                "Mobile Typing Friction:",
                "  • Native speakers face aggressive auto-correct errors defaulting to English/Hindi, forcing manual correction or total disabling of predictive features.",
                "Research Impact:",
                "  • Establishes the first digital NLP dataset, dictionary schema, and predictive input engine for Nagamese, bridging digital exclusion in Northeast India."
            ],
            "notes": "Slide 4 presents our Motivation under ML guideline item 3. Nagamese connects millions of speakers in Northeast India. However, smartphones lack native Nagamese dictionaries and language models. Users are forced to fight aggressive English auto-correct. Building native NLP models directly solves this digital divide and enables mobile digital inclusion."
        },
        {
            "num": "SLIDE 5",
            "title": "Literature Survey & Problem Statement (ML Guideline Items 4 & 5)",
            "bullets": [
                "1. LITERATURE SURVEY & GAPS (Item 4):",
                "   - Nagamese Linguistics (Sreedhar 1974, Baishya 2013, Boruah 2018): Documented Nagamese phonology and compounding. -> GAP: Purely descriptive academic studies with zero digital corpora or code artifacts.",
                "   - Low-Resource Indian NLP (Joshi et al. 2020): Highlighted 'Data Poverty' in regional languages. -> GAP: Northeastern creoles are absent from all Indic NLP benchmarks (IndicGLUE, Samanantar).",
                "   - Mobile Input Architecture (Fowler et al. 2015, Jurafsky & Martin 2023): Established Trie prefix indexing + N-gram backoff models for mobile typing. -> GAP: Zero implementations exist for Nagamese.",
                "2. PROBLEM STATEMENT (Item 5):",
                "   'Existing mobile platforms lack native language models and digital lexicons for Nagamese Creole, causing severe typing friction, aggressive auto-correct errors, and digital exclusion for Nagamese speakers.'"
            ],
            "notes": "Slide 5 covers the Literature Survey and Problem Statement under ML guideline items 4 and 5. Literature shows three major gaps: linguistic studies lack computational datasets; Indian NLP benchmarks ignore Northeastern creoles; and mobile predictive architectures have never been built for Nagamese. Our problem statement formally defines the typing friction and digital exclusion resulting from these gaps."
        },
        {
            "num": "SLIDE 6",
            "title": "Proposed Methodology – Architecture & Scope Split (ML Guideline Item 6a)",
            "bullets": [
                "TWO-TIER MODULAR SYSTEM ARCHITECTURE:",
                "  • Tier 1: Python NLP Research Pipeline (`nlp_research/`) -> Data preprocessing, Lexical DB compilation, N-Gram training, Trie indexing, and NMT prototyping.",
                "  • Tier 2: Android Demonstration Platform (`app/`) -> Native Kotlin IME keyboard host interface loading serialized model binaries offline.",
                "MINI PROJECT vs. MAJOR PROJECT SCOPE SPLIT:",
                "  ✅ MINI PROJECT (THIS REVIEW):",
                "     • Monolingual (6.9k lines) & Parallel Corpus (6.9k pairs)",
                "     • 20,000+ Entry Validated Lexical Database (`nagamese_lexicon.json`)",
                "     • N-Gram Language Model + Trie Prefix Tree Index",
                "     • Initial Functional Android IME Keyboard APK Demo",
                "  🔜 MAJOR PROJECT CONTINUATION:",
                "     • Neural Machine Translation (Seq2Seq / Transformer Nagamese <-> English)",
                "     • Final Polished APK Release & BLEU Score Evaluation"
            ],
            "notes": "Slide 6 presents our overall ML architecture and explicitly details the Mini vs. Major project scope split. The system uses a two-tier architecture separating Python model research from the Android keyboard app. For this Mini Project review, we deliver the complete dataset, 20k dictionary, N-gram model, Trie index, and initial Android keyboard APK. Neural Machine Translation (NMT) and final release polish are explicitly planned as the major project continuation."
        },
        {
            "num": "SLIDE 7",
            "title": "Proposed Methodology – Data Collection & Pre-processing (ML Stage 1 & 2)",
            "bullets": [
                "STAGE 1: DATA COLLECTION (ML Guideline Stage 1):",
                "  • Dataset Sources: 26 Nagamese scripture publications, digital web scrapers (xobdo.org, nagamesekhobor.com), curated code-switching loanwords list.",
                "  • Sample Size: 6,965 monolingual sentences (~185,945 running tokens), 6,965 aligned English-Nagamese parallel pairs.",
                "  • Data Type & Attributes: Text & Structured Lexical Schema (lemma, IPA, POS, English gloss, etymology).",
                "STAGE 2: DATA PRE-PROCESSING (ML Guideline Stage 2):",
                "  • Cleaning & Normalization: HTML/metadata removal, lowercase normalization across Romanized Nagamese text.",
                "  • Custom Nagamese Regex Tokenization: Isolated word tokens & sentence boundaries (`<s>`, `</s>`).",
                "  • Agglutinative Morphology Rules: Noun cases (`-khan`, `-laga`, `-ke`, `-pora`, `-te`) & Verb aspects (`-se`, `-bo`, `-bole`, `-ina`).",
                "  • Anti-Synthetic Purge Filter: Explicit rules blocking invalid compound tokens (e.g., *homolaga*)."
            ],
            "notes": "Slide 7 covers ML Pipeline Stages 1 and 2: Data Collection and Data Pre-processing. For Data Collection, we extracted 6,965 sentences from 26 scripture publications, scraped web glossaries, and curated 1,000+ code-switching loanwords. For Pre-processing, we implemented a custom regex tokenizer, morphological inflection rules, and an anti-synthetic filter that purges invalid compound words."
        },
        {
            "num": "SLIDE 8",
            "title": "Proposed Methodology – Feature Engineering & Model Selection (ML Stage 3, 4 & 5)",
            "bullets": [
                "STAGE 3: FEATURE ENGINEERING & EXTRACTION (ML Stage 3):",
                "  • Unigram, Bigram, and Trigram token transition counts.",
                "  • Character-level Trie node paths representing sub-word prefixes.",
                "  • Frequency-weighted candidate scoring for prediction ranking.",
                "STAGE 4: DATASET SPLITTING (ML Stage 4):",
                "  • 80% Training Set / 20% Evaluation Test Set for language model perplexity calculation.",
                "STAGE 5: MODEL SELECTION & RATIONALE (ML Stage 5):",
                "  • Selected Algorithm: Statistical N-Gram Language Model with Add-k Smoothing & Backoff + Character-Level Trie Index.",
                "  • Why Chosen Over Deep Learning (LSTMs/Transformers)?",
                "    1. High data efficiency on low-resource corpora (3,267 unique corpus tokens).",
                "    2. On-device mobile hardware constraints (<5 MB RAM footprint, sub-10ms response time).",
                "    3. 100% offline operation without cloud API dependencies."
            ],
            "notes": "Slide 8 details ML Pipeline Stages 3, 4, and 5: Feature Engineering, Dataset Splitting, and Model Selection. We extract N-gram transitions and Trie character paths as features. We use an 80/20 train/test split for perplexity evaluation. We selected Statistical N-Grams and Trie Trees over Deep Learning because N-grams require vastly less data, run offline under 5 MB RAM, and deliver sub-10 millisecond keyboard response times."
        },
        {
            "num": "SLIDE 9",
            "title": "Proposed Methodology – Training, Evaluation & Deployment (ML Stage 6, 7 & 8)",
            "bullets": [
                "STAGE 6: MODEL TRAINING (ML Stage 6):",
                "  • Add-k smoothing ($k=0.01$) to handle unseen word sequences.",
                "  • Context windowing (Trigram -> Bigram -> Unigram backoff).",
                "  • Model Size: 3,267 unigrams, 45,589 bigrams, 118,053 trigrams, 21,000 Trie dictionary nodes.",
                "STAGE 7: MODEL EVALUATION (ML Stage 7):",
                "  • Model Perplexity ($PP$): Achieved **45.59** (demonstrating high predictive certainty).",
                "  • Prediction Accuracy: Top-1, Top-3, and Top-5 candidate hit rates.",
                "  • Keystroke Savings (KS %): $\\text{KS} = \\left(1 - \\frac{\\text{Typed}}{\\text{Total Char}}\\right) \\times 100\\%$.",
                "STAGE 8: PREDICTION / DEPLOYMENT (ML Stage 8):",
                "  • Native Android `InputMethodService` keyboard application loading `trie_index.json` (0.77 MB) & `bigrams.json` offline.",
                "  • Real-time candidate bar rendering (<10 ms latency, <5 MB RAM)."
            ],
            "notes": "Slide 9 covers ML Pipeline Stages 6, 7, and 8: Training, Evaluation, and Deployment. We trained N-gram models with add-k smoothing, achieving a perplexity of 45.59. We evaluate model performance using Perplexity, Top-k accuracy, and Keystroke Savings. For Deployment, the models are serialized into compact JSON assets loaded offline by our native Android keyboard APK."
        },
        {
            "num": "SLIDE 10",
            "title": "Step-by-Step Workflow & References (ML Guideline Item 6c & 6d)",
            "bullets": [
                "STEP-BY-STEP ML WORKFLOW (Item 6c):",
                "  1. Collect dataset (26 PDFs + web scrapers + loanwords).",
                "  2. Preprocess & clean text (regex tokenization + morphological filtering).",
                "  3. Extract N-gram transition features & build Trie node paths.",
                "  4. Split dataset into training (80%) and evaluation (20%) sets.",
                "  5. Train statistical N-gram language model with add-k smoothing.",
                "  6. Evaluate model perplexity ($PP = 45.59$) and Top-k accuracy.",
                "  7. Export compact model binaries (`trie_index.json`, `bigrams.json`).",
                "  8. Deploy on Android IME platform and validate offline keystroke autocompletion.",
                "KEY REFERENCES:",
                "  • Sreedhar (1974) - Nagamese | Baishya (2013) - Compounding | Joshi et al. (2020) - Low-Resource NLP",
                "  • Fowler et al. (2015) - Trie Mobile Input | Jurafsky & Martin (2023) - N-gram Language Modeling"
            ],
            "notes": "Slide 10 presents the Step-by-Step ML Workflow following guideline item 6c, summarizing our 8-step sequence from data collection to mobile deployment. It also highlights our key academic references. Thank you very much, Nzanthung sir, Nokshangthemba sir, and evaluation committee members. I am open for your questions."
        }
    ]

    for slide in slides:
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(3)
        r_num = h2.add_run(f"{slide['num']}: ")
        r_num.bold = True
        r_num.font.size = Pt(13)
        r_num.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
        r_title = h2.add_run(slide['title'])
        r_title.bold = True
        r_title.font.size = Pt(13)
        r_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        for bullet in slide['bullets']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.size = Pt(9.5)

        add_callout_box(doc, "Speaker Notes (What to say)", slide['notes'])

    # Defense Q&A Section
    doc.add_page_break()
    hq_p = doc.add_paragraph()
    hq_p.paragraph_format.space_before = Pt(12)
    hq_p.paragraph_format.space_after = Pt(8)
    hq_r = hq_p.add_run("DEFENSE Q&A — MINI PROJECT ML PIPELINE SPECIFIC")
    hq_r.bold = True
    hq_r.font.size = Pt(15)
    hq_r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    qas = [
        ("Q1: How does your NLP project align with the ML/AI pipeline specified in the project guidelines?",
         "Answer: Sir, NLP is a core subfield of AI/ML handling text data. My project strictly follows the ML pipeline specified in the guidelines: (1) Data Collection: 6.9k sentences from PDFs/web scraping; (2) Data Pre-processing: regex tokenization, morphological filtering, anti-synthetic purge; (3) Feature Extraction: N-gram transition probabilities & Trie character paths; (4) Model Selection: Statistical N-Grams with Add-k Smoothing & Trie Index; (5) Model Evaluation: Perplexity (45.59) & Keystroke Savings; (6) Deployment: On-device Android IME application."),
        ("Q2: Why is Neural Machine Translation (NMT) deferred to the Major Project phase?",
         "Answer: Neural Machine Translation using seq2seq or Transformer architectures requires extensive hyperparameter tuning, larger parallel training data, and complex BLEU evaluation. For the Mini Project (Review-I), our goal is to build foundational computational resources (corpus + 20k dictionary) and prove practical viability through the predictive keyboard engine. NMT is the natural research continuation for the major project."),
        ("Q3: Why choose Statistical N-Grams over Deep Learning (LSTMs/Transformers) for word prediction?",
         "Answer: There are two engineering reasons: First, deep neural models require millions of training sentences, whereas Nagamese is a low-resource creole. Second, our target platform is an on-device mobile keyboard that must operate offline with <5 MB RAM footprint and <10ms latency. Statistical N-grams combined with Trie prefix trees achieve superior efficiency and speed on mobile hardware."),
        ("Q4: Is the Android keyboard APK delivered as part of the Mini Project?",
         "Answer: Yes, sir. The Android keyboard APK is delivered as the Mini Project demonstration platform. It loads the serialized model binaries (trie_index.json and bigrams.json) offline and renders real-time word suggestions in the candidate bar as the user types.")
    ]

    for q, a in qas:
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        qrun = qp.add_run(q)
        qrun.bold = True
        qrun.font.size = Pt(11)
        qrun.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

        ap = doc.add_paragraph()
        ap.paragraph_format.space_after = Pt(8)
        arun = ap.add_run(a)
        arun.font.size = Pt(10.5)

    doc.save(output_path)
    print(f"Successfully generated Fresh ML-Aligned Mini Project DOCX at: {output_path}")

if __name__ == "__main__":
    out_path = "f:/likhibi-main/docs/Nagamese_NLP_Mini_Project_Presentation.docx"
    build_mini_project_docx(out_path)
