import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, title, text, bg_hex="F0F4F8", border_hex="1A365D"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"🗣️ {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.italic = True
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def create_10_slide_docx(output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    styles = doc.styles
    normal_font = styles['Normal'].font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("PROJECT REVIEW – I (10-SLIDE STREAMLINED SCRIPT)")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run("Major Project Proposal Slide Deck Script & Concise Speaker Notes\nLikhibi: Nagamese NLP Resources & Predictive Input System")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Metadata Box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Domain:", "Natural Language Processing (NLP) / Computational Linguistics"),
        ("Project Coordinators:", "Mr. Nzanthung Odyuo | Mr. Nokshangthemba")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width, cell_v.width = Inches(2.0), Inches(4.5)
        set_cell_background(cell_k, "F7FAFC")
        set_cell_background(cell_v, "FFFFFF")
        
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(2)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(9.5)
        rk.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(2)
        rv = pv.add_run(v)
        rv.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    slides_10 = [
        {
            "num": "SLIDE 1",
            "title": "Title Slide",
            "bullets": [
                "Title: Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole",
                "Review: Major Project Proposal Defense (Project Review – I)",
                "Domain: Natural Language Processing / Computational Linguistics",
                "Target Language: Nagamese Creole (Lingua Franca of Nagaland)",
                "Presenter: [Your Name] | Roll No / USN: [Your Roll Number] | Dept of CSE",
                "Project Coordinators: Mr. Nzanthung Odyuo & Mr. Nokshangthemba"
            ],
            "notes": "Respected Project Coordinators Mr. Nzanthung Odyuo sir, Mr. Nokshangthemba sir, and evaluation committee members. I am [Your Name], presenting my project proposal: 'Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole'. Nagamese is the primary spoken lingua franca across Nagaland, but remains a Low-Resource Language in computer science. This project builds native computational NLP resources and an offline predictive text engine for mobile devices."
        },
        {
            "num": "SLIDE 2",
            "title": "Table of Contents",
            "bullets": [
                "1. Project Core (Aim, Objectives & Motivation)",
                "2. Literature Survey & Gap Analysis",
                "3. Problem Statement",
                "4. Proposed Methodology: Architecture & Pipeline",
                "5. Data Acquisition, Preprocessing & Lexical Database",
                "6. N-Gram Language Modeling & Trie Indexing",
                "7. Step-by-Step Workflow & Mobile IME Deployment",
                "8. Expected Outcomes, Evaluation Metrics & References"
            ],
            "notes": "Here is the 10-slide outline for today's proposal review. I will cover the core project parameters, survey existing literature, define the problem statement, and detail our end-to-end NLP methodology from data collection to mobile deployment."
        },
        {
            "num": "SLIDE 3",
            "title": "Aim, Objectives & Motivation (Core Summary)",
            "bullets": [
                "AIM: Create standardized NLP computational resources for Nagamese Creole and deploy an offline contextual word-prediction engine integrated into an Android Keyboard (IME).",
                "KEY OBJECTIVES:",
                "  • Corpus Building: Construct monolingual corpus (~7,000 lines) & parallel corpus.",
                "  • Lexical Database: Build a verified 20,000+ entry Nagamese digital dictionary.",
                "  • Language Models: Train statistical N-gram models (Unigram, Bigram, Trigram) with add-k smoothing.",
                "  • Trie Indexing: Construct character-level Trie for sub-millisecond prefix completion.",
                "  • Mobile Deployment: Integrate models into an offline Android IME (<5 MB RAM).",
                "MOTIVATION:",
                "  • Spoken by 30M+ people, yet ZERO native mobile keyboard support exists on Android/iOS in Roman script.",
                "  • Eliminates aggressive English auto-correct errors for Nagamese users & bridges regional NLP data poverty."
            ],
            "notes": "Slide 3 combines our core foundations: Our Aim is to build digital NLP resources for Nagamese and deploy an offline predictive keyboard. Our Objectives focus on 5 concrete milestones: corpus creation, a 20,000-entry dictionary, N-gram language models, a Trie index, and Android keyboard integration. Our Motivation stems from sociolinguistic reality: Nagamese is spoken by millions daily, yet smartphones lack any native prediction, forcing users to fight English auto-correct."
        },
        {
            "num": "SLIDE 4",
            "title": "Literature Survey & Gap Analysis",
            "bullets": [
                "1. Nagamese Linguistic Foundations (Sreedhar 1974, Baishya 2013, Boruah 2018): Documented Nagamese grammar, noun/verb compounding. -> GAP: Purely descriptive academic studies without digital corpora, tokenizers, or code artifacts.",
                "2. Indian Low-Resource NLP (Joshi et al. 2020, Kumar et al. 2021): Highlighted severe 'Data Poverty' in regional Indian languages. -> GAP: Northeastern creoles are completely missing from Indic NLP benchmarks (e.g., IndicGLUE).",
                "3. Mobile Predictive Input Architecture (Fowler et al. 2015, Jurafsky & Martin 2023): Established Trie prefix indexing + N-gram backoff models for mobile typing. -> GAP: Zero predictive input models or keyboard engines exist for Nagamese."
            ],
            "notes": "Literature shows three major gaps: First, linguistic studies on Nagamese are purely academic with no digital datasets. Second, Indian NLP benchmarks completely ignore Northeastern creoles. Third, mobile predictive input research has never been applied to Nagamese. Our project directly bridges these three gaps."
        },
        {
            "num": "SLIDE 5",
            "title": "Problem Statement",
            "bullets": [
                "Problem Definition: Existing mobile operating systems lack native language models and digital lexicons for Nagamese Creole, causing severe typing friction, aggressive auto-correct errors, and digital exclusion for Nagamese speakers.",
                "Broader Impact: The total absence of standardized digital corpora prevents the development of downstream NLP technologies (such as Machine Translation) for the language."
            ],
            "notes": "The problem statement is straightforward: Nagamese speakers face constant typing friction because mobile operating systems lack native dictionaries and language models. Additionally, researchers cannot build machine translation tools due to the lack of standardized digital corpora."
        },
        {
            "num": "SLIDE 6",
            "title": "Proposed Methodology – Architecture Overview",
            "bullets": [
                "Stage 1: Data Acquisition (PDF Scripture Extraction + Digital Web Scraping)",
                "Stage 2: Text Preprocessing (Cleaner + Custom Tokenizer + Normalizer)",
                "Stage 3: Resource Creation (20,000+ Entry Lexical DB + Monolingual/Parallel Corpora)",
                "Stage 4: Model Building (Trie Prefix Tree Indexer + Statistical N-Gram Trainer)",
                "Stage 5: Hybrid Prediction Engine (Reranking & Context Matching) -> On-Device Android IME Keyboard App"
            ],
            "notes": "Slide 6 shows our complete engineering pipeline divided into 5 stages: Data Acquisition, Text Preprocessing, Resource Creation (Lexicon + Corpus), Model Building (Trie + N-Gram), and Android IME Deployment."
        },
        {
            "num": "SLIDE 7",
            "title": "Proposed Methodology – Data, Preprocessing & Lexicon",
            "bullets": [
                "Custom Data Acquisition:",
                "  • Monolingual Corpus: 26 Nagamese scripture publications (6,965 sentences, ~185k tokens).",
                "  • Scraped Data: Web glossaries and community lexicons (xobdo.org, nagamesekhobor.com).",
                "  • Parallel Corpus: 6,965 aligned English-Nagamese sentence pairs (.tsv).",
                "  • Code-Switching Dataset: 1,000+ high-frequency English/Hindi loanwords (school, phone, office, time).",
                "Preprocessing & Morphological Filter:",
                "  • Custom regex tokenizer tailored for Romanized creole text with boundary markers (<s>, </s>).",
                "  • Creole morphology inflections: Noun cases (-khan, -laga, -ke, -pora, -te) & Verb aspects (-se, -bo, -bole, -ina).",
                "  • Anti-Synthetic Purge Filter: Explicit rules blocking invalid compounds (e.g., homolaga).",
                "20,000+ Entry Lexical Database (nagamese_lexicon.json):",
                "  • Schema: Lemma, IPA phonetics, POS category, English definition, frequency, etymology tag.",
                "  • Ratio: 95.2% Native Nagamese + 3.8% English Loanwords + 1.0% Hindi Borrowings. 100% automated validation scan."
            ],
            "notes": "On Slide 7, we detail the data engine: We built a custom 6,965-line monolingual corpus, a 6,965-pair parallel corpus, and a 1,000-word code-switching dataset. Preprocessing includes custom regex tokenization and creole morphological rules with an anti-synthetic filter. The output is a 20,000+ entry verified dictionary in JSON schema with etymology tagging."
        },
        {
            "num": "SLIDE 8",
            "title": "Proposed Methodology – N-Gram & Trie Indexing",
            "bullets": [
                "Statistical N-Gram Language Model:",
                "  • Trains Unigram, Bigram, and Trigram probabilities with Add-k Smoothing: P(w_i | w_{i-1}) = (C(w_{i-1}, w_i) + k) / (C(w_{i-1}) + k * |V|).",
                "  • Uses Backoff smoothing for unseen sequences (Trigram -> Bigram -> Unigram).",
                "  • Model Perplexity Metric (PP): Evaluates predictive certainty on test sequences (Target: PP < 50).",
                "Character-Level Trie Prefix Tree:",
                "  • Character tree indexing 20,000+ dictionary entries with O(L) search latency (<1 ms).",
                "  • Serialized to compact JSON payload (~0.77 MB).",
                "Hybrid Reranking Engine:",
                "  • User typing prefix ('ja') -> Trie matches candidates ['jabo', 'jai', 'jani'].",
                "  • Context exists ('moi') -> N-gram context probability boosts ranking: FinalScore(w) = TrieFreq(w) + alpha * P_Ngram(w | Context)."
            ],
            "notes": "Slide 8 details our prediction algorithms: We train N-gram models with Add-k smoothing and backoff to predict next words given context. We build a Trie prefix tree for O(L) sub-millisecond prefix completion. A hybrid reranker combines Trie candidate frequencies with N-gram context probabilities to surface top predictions."
        },
        {
            "num": "SLIDE 9",
            "title": "Workflow & Mobile IME Deployment",
            "bullets": [
                "Android IME Architecture:",
                "  • Built on native Android InputMethodService API.",
                "  • Storage Footprint: Serialized trie_index.json & bigrams.json loaded into RAM (<5 MB total).",
                "  • Processing Flow: User Type -> Capture Prefix & Context -> Trie/N-Gram Query -> Render Top-3/Top-5 Suggestions.",
                "  • Core Constraints: 100% Offline operation, <10 ms response latency.",
                "Step-by-Step Project Workflow:",
                "  • Phase I: Scripture extraction & text tokenization.",
                "  • Phase II: Lexical database curation, stem inflections & 100% validation scan.",
                "  • Phase III: N-Gram model training, Trie serialization & Perplexity evaluation.",
                "  • Phase IV: Android IME UI integration & Keystroke Savings evaluation."
            ],
            "notes": "Slide 9 covers execution and deployment: The engine is integrated into an Android mobile keyboard using InputMethodService. The model requires under 5 MB RAM and runs 100% offline with zero network latency. The workflow spans 4 structured phases from corpus extraction to mobile IME evaluation."
        },
        {
            "num": "SLIDE 10",
            "title": "Expected Outcomes, Metrics & References",
            "bullets": [
                "Deliverables: (1) Validated 20,000+ Dictionary nagamese_lexicon.json, (2) Parallel Corpus bible_parallel_corpus.tsv (6,965 pairs), (3) Serialized Model Binaries trie_index.json & bigrams.json, (4) Working Offline Android Keyboard APK.",
                "Evaluation Metrics:",
                "  • Perplexity (PP): Measures prediction uncertainty (Target: PP < 50).",
                "  • Keystroke Savings (KS %): KS = (1 - Typed_Keystrokes / Total_Characters) * 100%.",
                "  • Accuracy: Top-1, Top-3, and Top-5 suggestion hit rates.",
                "  • Hardware Performance: Latency (<10 ms) and Memory Footprint (<10 MB RAM).",
                "Key References:",
                "  • Sreedhar (1974) - Nagamese Linguistics | Baishya (2013) - Nagamese Compounding",
                "  • Joshi et al. (2020) - Low-Resource Indian NLP | Fowler et al. (2015) - Trie Mobile IME"
            ],
            "notes": "Finally, Slide 10 summarizes deliverables and evaluation metrics: Key outcomes include the 20k dictionary, parallel corpus, serialized models, and Android APK. Success is measured via Perplexity, Keystroke Savings percentage, Top-k accuracy, and mobile hardware benchmarks. Thank you, Nzanthung sir, Nokshangthemba sir, and committee members. I am ready for your questions."
        }
    ]

    for slide in slides_10:
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(3)
        
        r_num = h2.add_run(f"{slide['num']}: ")
        r_num.bold = True
        r_num.font.size = Pt(13)
        r_num.font.color.rgb = RGBColor(0x31, 0x82, 0xCE)
        
        r_title = h2.add_run(slide['title'])
        r_title.bold = True
        r_title.font.size = Pt(13)
        r_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        for bullet in slide['bullets']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.size = Pt(9.5)

        add_callout_box(doc, "Speaker Notes (What to say)", slide['notes'])

    doc.save(output_path)
    print(f"Successfully created 10-Slide DOCX at: {output_path}")

if __name__ == "__main__":
    out = "f:/likhibi-main/docs/Nagamese_NLP_10_Slide_Presentation.docx"
    create_10_slide_docx(out)
