import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, title, text, bg_hex="F0F4F8", border_hex="1A365D"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"🗣️ {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.italic = True
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def create_mini_project_docx(output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = doc.styles
    normal_font = styles['Normal'].font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Document Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("MINI PROJECT REVIEW — PROPOSAL PRESENTATION SCRIPT")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run("10-Slide Proposal Script & Speaker Notes (Mini Project Scope)\nLikhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Review Type:", "Mini Project Proposal Defense (Project Review – I)"),
        ("Mini Project Scope:", "Corpus + 20k Lexicon + N-Gram + Trie Prediction + Initial Android IME APK"),
        ("Project Coordinators:", "Mr. Nzanthung Odyuo | Mr. Nokshangthemba")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width, cell_v.width = Inches(2.0), Inches(4.5)
        set_cell_background(cell_k, "F7FAFC")
        set_cell_background(cell_v, "FFFFFF")
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(2)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(9.5)
        rk.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(2)
        rv = pv.add_run(v)
        rv.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Scope Split Box
    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_table.autofit = False
    mini_cell = scope_table.cell(0, 0)
    major_cell = scope_table.cell(0, 1)
    mini_cell.width = Inches(3.2)
    major_cell.width = Inches(3.3)

    set_cell_background(mini_cell, "EBF8FF")
    set_cell_margins(mini_cell, top=120, bottom=120, left=150, right=100)
    tcPr_m = mini_cell._element.get_or_add_tcPr()
    borders_m = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2196F3"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr_m.append(borders_m)
    p_m = mini_cell.paragraphs[0]
    r_m1 = p_m.add_run("✅ MINI PROJECT DELIVERABLES\n")
    r_m1.bold = True; r_m1.font.size = Pt(9.5); r_m1.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    r_m2 = p_m.add_run("• Monolingual & Parallel Corpus\n• 20,000+ Entry Lexical Database\n• N-Gram Language Model\n• Trie Prefix Prediction Engine\n• Initial Android IME Keyboard APK")
    r_m2.font.size = Pt(9.5)

    set_cell_background(major_cell, "FFF3E0")
    set_cell_margins(major_cell, top=120, bottom=120, left=150, right=100)
    tcPr_ma = major_cell._element.get_or_add_tcPr()
    borders_ma = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="F57C00"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr_ma.append(borders_ma)
    p_ma = major_cell.paragraphs[0]
    r_ma1 = p_ma.add_run("🔜 MAJOR PROJECT (CONTINUATION)\n")
    r_ma1.bold = True; r_ma1.font.size = Pt(9.5); r_ma1.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    r_ma2 = p_ma.add_run("• Neural Machine Translation (NMT)\n• Nagamese ↔ English seq2seq model\n• Final polished APK release\n• Complete research documentation\n• BLEU score evaluation")
    r_ma2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ─── SLIDES ───
    slides = [
        {
            "num": "SLIDE 1",
            "title": "Title Slide",
            "bullets": [
                "Project Title: Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole",
                "Review: Mini Project Proposal Defense (Project Review – I)",
                "Domain: Natural Language Processing / Computational Linguistics",
                "Target Language: Nagamese Creole (Lingua Franca of Nagaland)",
                "Presenter: [Your Name] | Roll No / USN: [Your Roll Number] | Dept of CSE",
                "Project Coordinators: Mr. Nzanthung Odyuo & Mr. Nokshangthemba"
            ],
            "notes": "Respected Project Coordinators Mr. Nzanthung Odyuo sir, Mr. Nokshangthemba sir, and evaluation committee members. I am [Your Name], presenting my mini project proposal: 'Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole'. Nagamese is the primary spoken lingua franca of Nagaland with over 30 million daily users, yet it remains a completely Low-Resource Language in computing with no native digital keyboard support. This mini project builds the foundational NLP resources and delivers a functional initial Android keyboard as the demonstration platform."
        },
        {
            "num": "SLIDE 2",
            "title": "Table of Contents",
            "bullets": [
                "1. Aim, Objectives & Motivation",
                "2. Literature Survey & Gap Analysis",
                "3. Problem Statement",
                "4. Mini vs. Major Project Scope Split",
                "5. Proposed Methodology Overview & System Architecture",
                "6. Data Acquisition, Preprocessing & 20,000 Lexical Database",
                "7. N-Gram Language Modeling & Trie Prefix Indexing",
                "8. Android IME Keyboard – Initial APK (Mini Project Demo Platform)",
                "9. Step-by-Step Workflow & Timeline",
                "10. Expected Outcomes, Evaluation Metrics & References"
            ],
            "notes": "Here is the 10-slide outline for today's mini project proposal review. I will cover our aims, survey literature gaps, define the problem, explain the mini vs. major project split, then walk through our complete data and NLP methodology—finishing with the initial Android keyboard as the mini project demonstration deliverable."
        },
        {
            "num": "SLIDE 3",
            "title": "Aim, Objectives & Motivation (Combined)",
            "bullets": [
                "AIM: Build foundational NLP computational resources for Nagamese Creole and deploy an offline contextual word-prediction engine integrated into an initial Android IME keyboard.",
                "MINI PROJECT OBJECTIVES (THIS REVIEW):",
                "  • Corpus: Construct 6,965-line monolingual corpus & 6,965-pair English-Nagamese parallel corpus.",
                "  • Lexical Database: Build a verified 20,000+ entry Nagamese dictionary (nagamese_lexicon.json).",
                "  • N-Gram Language Model: Train Unigram, Bigram, Trigram models with add-k smoothing.",
                "  • Trie Prefix Tree: Character-level index for sub-millisecond prefix completion.",
                "  • Initial Android IME: Functional keyboard APK integrating prediction engine offline.",
                "MOTIVATION:",
                "  • 30M+ speakers, yet ZERO native mobile keyboard or prediction support in Roman script.",
                "  • Bridges digital exclusion for Nagamese and demonstrates research viability for major continuation."
            ],
            "notes": "Slide 3 covers our combined project foundations. The aim of the mini project is to build the core data resources and deploy a working predictive keyboard. We have 5 concrete mini project objectives: building a corpus, 20k dictionary, N-gram model, Trie index, and initial functional Android IME APK. The motivation is simple: 30 million people speak Nagamese daily but cannot type naturally on smartphones due to missing native support."
        },
        {
            "num": "SLIDE 4",
            "title": "Literature Survey & Gap Analysis",
            "bullets": [
                "1. Nagamese Linguistic Foundations (Sreedhar 1974, Baishya 2013, Boruah 2018): Documented Nagamese grammar, creole compounding, and syntax. -> GAP: Zero computational datasets, tokenizers, or digital corpora exist.",
                "2. Indian Low-Resource NLP (Joshi et al. 2020): Highlighted 'Data Poverty' in Indian regional languages. -> GAP: Northeastern creoles are absent from all Indic NLP benchmarks (IndicGLUE, Samanantar).",
                "3. Mobile Predictive Input Architecture (Fowler et al. 2015, Jurafsky & Martin 2023): Trie + N-gram backoff optimal for mobile IME. -> GAP: No Trie/N-gram implementation or IME exists for Nagamese.",
                "Gaps Addressed by This Mini Project: (1) First validated 20k digital dictionary, (2) First trained N-gram language model, (3) First Trie prefix index, (4) First functional Nagamese Android keyboard APK."
            ],
            "notes": "Literature review shows three major gaps: First, linguistic studies on Nagamese are purely academic with no code artifacts. Second, Indian NLP benchmarks completely ignore Northeastern creoles. Third, mobile keyboard input methods using Trie and N-gram have never been implemented for Nagamese. Our mini project addresses all four of these gaps directly."
        },
        {
            "num": "SLIDE 5",
            "title": "Problem Statement",
            "bullets": [
                "Problem: Existing mobile operating systems lack native language models and digital lexicons for Nagamese Creole, causing aggressive auto-correct errors, typing friction, and digital exclusion for Nagamese speakers.",
                "Impact 1: Native speakers are forced to type in incorrect English or disable auto-correct entirely.",
                "Impact 2: The absence of standardized digital corpora blocks downstream NLP research (translation, sentiment analysis) for the language.",
                "This Mini Project Addresses Impact 1 directly by delivering a functional offline predictive IME keyboard."
            ],
            "notes": "The problem is straightforward: Nagamese speakers face constant typing friction because mobile systems lack native dictionaries and language models. This mini project directly solves Impact 1 by delivering a working Android keyboard with offline prediction. Impact 2—the research corpus gap—is also addressed through our corpus and lexical database, enabling future work on machine translation in the major project continuation."
        },
        {
            "num": "SLIDE 6",
            "title": "Mini vs. Major Project Scope Split",
            "bullets": [
                "MINI PROJECT (Current Review — Submitted This Year):",
                "  ✅ Corpus Creation: 6,965-line monolingual + 6,965-pair parallel corpus",
                "  ✅ 20,000+ Entry Verified Lexical Database (nagamese_lexicon.json)",
                "  ✅ N-Gram Language Model (Unigram, Bigram, Trigram + Add-k Smoothing)",
                "  ✅ Character-Level Trie Prefix Index (0.77 MB serialized JSON)",
                "  ✅ Initial Android IME Keyboard APK (offline, <5 MB RAM, <10 ms response)",
                "",
                "MAJOR PROJECT CONTINUATION (Next Phase):",
                "  🔜 Neural Machine Translation (Seq2Seq / Transformer, Nagamese <-> English)",
                "  🔜 Final Polished Production-Ready Android APK Release",
                "  🔜 BLEU Score Evaluation & Full Research Documentation",
                "  🔜 Publication-Ready NLP Resource Paper"
            ],
            "notes": "This slide clearly defines our scope split for the committee. The mini project is complete in scope: corpus, 20k dictionary, N-gram model, Trie index, and a functional initial keyboard APK. These are concrete, testable deliverables. The major project continuation adds Neural Machine Translation—which requires more training data and compute—along with a finalized polished release and academic documentation."
        },
        {
            "num": "SLIDE 7",
            "title": "Proposed Methodology – Data & Lexical Database",
            "bullets": [
                "Data Collection:",
                "  • Monolingual Corpus: Extracted from 26 Nagamese scripture publications (6,965 sentences, ~185k tokens).",
                "  • Web Scraped: Regional glossaries and community lexicons (xobdo.org, nagamesekhobor.com) -> 1,500+ contemporary tokens.",
                "  • Parallel Corpus: 6,965 aligned English-Nagamese sentence pairs (bible_parallel_corpus.tsv).",
                "  • Code-Switching Dataset: 1,000+ frequent English/Hindi loanwords used in Nagamese code-switching.",
                "Preprocessing Pipeline:",
                "  • Custom regex tokenizer for Romanized creole text with sentence boundary markers (<s>, </s>).",
                "  • Morphological inflections: Noun cases (-khan, -laga, -ke, -pora, -te) & Verb aspects (-se, -bo, -bole, -ina).",
                "  • Anti-Synthetic Purge Filter: Blocks invalid compound tokens (e.g., homolaga).",
                "20,000+ Entry Lexical Database:",
                "  • Schema: lemma, IPA, POS, English definition, frequency, etymology. 100% automated validation scan. 0 invalid entries."
            ],
            "notes": "On Slide 7, we detail the data pipeline: custom corpus extraction from 26 scripture publications, web scraping for contemporary vocabulary, and a 6,965-pair parallel corpus. The preprocessing uses a custom tokenizer with morphological rules specific to Nagamese creole structure. The output is a 20,000+ entry dictionary verified against primary sources with zero invalid entries."
        },
        {
            "num": "SLIDE 8",
            "title": "Proposed Methodology – N-Gram & Trie Prediction",
            "bullets": [
                "N-Gram Language Model (ngram_model.py):",
                "  • Unigram, Bigram, Trigram frequency tables with Add-k Smoothing: P(w_i | w_{i-1}) = (C(w_{i-1}, w_i) + k) / (C(w_{i-1}) + k * |V|).",
                "  • Backoff: Trigram -> Bigram -> Unigram for unseen sequences.",
                "  • Corpus Stats: 3,267 vocabulary, 45,589 bigrams, 118,053 trigrams. Perplexity: 45.59.",
                "  • Export: bigrams.json, trigrams.json for Android asset loading.",
                "Character-Level Trie Prefix Index (trie_builder.py):",
                "  • Indexes all 20,000+ lemmas in character tree structure.",
                "  • O(L) search time (sub-1 ms). Serialized to trie_index.json (0.77 MB).",
                "Hybrid Reranking Engine (prediction_engine.py):",
                "  • Prefix 'ja' -> Trie: ['jabo', 'jai', 'jani']. Context 'moi' -> N-gram reranks by probability.",
                "  • FinalScore(w) = TrieFreq(w) + alpha * P_Ngram(w | Context)."
            ],
            "notes": "Slide 8 details the prediction algorithms. The N-gram model was trained on our 6,965-sentence corpus and achieves a perplexity of 45.59 with add-k smoothing. The Trie prefix index covers all 20,000 words and executes in under 1 millisecond. The hybrid engine combines both: Trie candidate frequencies are reranked by N-gram contextual probability to put the most likely next word first."
        },
        {
            "num": "SLIDE 9",
            "title": "Initial Android IME Keyboard (Mini Project Demo)",
            "bullets": [
                "Platform: Native Android InputMethodService (Kotlin).",
                "Integration Architecture:",
                "  • trie_index.json + bigrams.json bundled as Android assets (<5 MB total).",
                "  • PredictionEngine interface connects keyboard input to model queries.",
                "  • Top-3 / Top-5 suggestions rendered in candidate bar above keyboard.",
                "Key Performance Targets:",
                "  • 100% Offline operation (no internet required).",
                "  • <10 ms prediction response time.",
                "  • <5 MB RAM footprint.",
                "Mini Project Timeline (Phases I - III):",
                "  • Phase I: Corpus extraction, preprocessing & parallel alignment.",
                "  • Phase II: 20k lexical database curation & 100% validation scan.",
                "  • Phase III: N-Gram training, Trie building, export & Android IME APK integration."
            ],
            "notes": "Slide 9 covers the mini project's primary demonstration deliverable—the initial Android IME keyboard. The model assets are bundled directly into the APK as Android assets, keeping the app fully offline. The prediction engine interfaces with the keyboard's candidate bar to show suggestions in real time. The mini project implementation spans 3 clear phases: corpus, database, and model plus Android integration."
        },
        {
            "num": "SLIDE 10",
            "title": "Expected Outcomes, Metrics & References",
            "bullets": [
                "Mini Project Deliverables:",
                "  • Validated 20,000+ Entry Nagamese Lexical Database (nagamese_lexicon.json, 0 invalid entries).",
                "  • Monolingual Corpus (6,965 lines) + Parallel Corpus (6,965 pairs, bible_parallel_corpus.tsv).",
                "  • Trained N-Gram Model Assets (unigrams.json, bigrams.json, trigrams.json).",
                "  • Trie Prefix Index (trie_index.json, 0.77 MB).",
                "  • Initial Functional Android IME Keyboard APK.",
                "Evaluation Metrics (Mini Project):",
                "  • Model Perplexity (PP): N-Gram predictive certainty [Achieved: 45.59].",
                "  • Prediction Accuracy: Top-1, Top-3, Top-5 suggestion accuracy rates.",
                "  • Keystroke Savings (KS %): KS = (1 - Typed / Total_Chars) * 100%.",
                "  • Hardware: <10 ms latency, <5 MB RAM footprint.",
                "Key References:",
                "  • Sreedhar (1974) - Nagamese | Joshi et al. (2020) - Low-Resource NLP | Fowler et al. (2015) - Trie IME"
            ],
            "notes": "Finally, Slide 10 summarizes our mini project deliverables and evaluation metrics. The key outcomes are the 20k dictionary with zero invalid entries, the trained N-gram model with perplexity 45.59, the 0.77 MB Trie index, and the initial Android keyboard APK. We evaluate using Perplexity, Keystroke Savings percentage, and Top-k prediction accuracy. Thank you, Nzanthung sir, Nokshangthemba sir, and committee members. I am ready for your questions."
        }
    ]

    for slide in slides:
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(3)
        r_num = h2.add_run(f"{slide['num']}: ")
        r_num.bold = True
        r_num.font.size = Pt(13)
        r_num.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
        r_title = h2.add_run(slide['title'])
        r_title.bold = True
        r_title.font.size = Pt(13)
        r_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        for bullet in slide['bullets']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.size = Pt(9.5)

        add_callout_box(doc, "Speaker Notes (What to say)", slide['notes'])

    # Defense Q&A
    doc.add_page_break()
    hq_p = doc.add_paragraph()
    hq_p.paragraph_format.space_before = Pt(12)
    hq_p.paragraph_format.space_after = Pt(8)
    hq_r = hq_p.add_run("DEFENSE Q&A — MINI PROJECT SPECIFIC")
    hq_r.bold = True
    hq_r.font.size = Pt(15)
    hq_r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    qas = [
        ("Q1: Why is Neural Machine Translation not part of the mini project?",
         "Answer: Machine Translation using neural seq2seq models requires substantially more training data and compute time to produce meaningful results. With only 6,965 parallel sentence pairs, a baseline NMT model needs careful hyperparameter tuning and evaluation—work best suited for the extended major project phase. The mini project focuses on the core resources and demonstrates practical value through the predictive keyboard."),
        ("Q2: Is the initial Android APK fully functional?",
         "Answer: Yes. The initial APK integrates the N-gram model and Trie prediction engine offline via bundled JSON asset files. The keyboard renders Top-3 to Top-5 suggestions in the candidate bar in under 10 milliseconds with no internet required. What is deferred to the major project is the final polished UI/UX refinement, Play Store release packaging, and the translation feature layer."),
        ("Q3: How is this a mini project when you have 20,000 dictionary entries and trained models?",
         "Answer: The scale of resources we built actually strengthens the mini project—it shows thorough work within the defined scope. The mini project scope is complete: corpus, dictionary, N-gram model, Trie, and an initial keyboard. The continuation to major project adds an entirely new component—Neural Machine Translation—which represents a separate research challenge requiring different algorithms, datasets, and evaluation methods."),
        ("Q4: Why N-grams instead of deep learning for prediction?",
         "Answer: N-gram models with Trie indexing execute in under 2 milliseconds offline on mobile hardware with less than 5 MB RAM—requirements that deep learning models cannot currently meet for on-device keyboards. Furthermore, the 6,965-sentence corpus is too small to train reliable neural language models for prediction. Statistical methods are the research-appropriate and deployment-practical choice for this phase.")
    ]

    for q, a in qas:
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        qrun = qp.add_run(q)
        qrun.bold = True
        qrun.font.size = Pt(11)
        qrun.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

        ap = doc.add_paragraph()
        ap.paragraph_format.space_after = Pt(8)
        arun = ap.add_run(a)
        arun.font.size = Pt(10.5)

    doc.save(output_path)
    print(f"Successfully created Mini Project DOCX at: {output_path}")

if __name__ == "__main__":
    out = "f:/likhibi-main/docs/Nagamese_NLP_Mini_Project_Proposal.docx"
    create_mini_project_docx(out)
