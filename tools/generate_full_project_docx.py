import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, title, text, bg_hex="F0F4F8", border_hex="1A365D"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"🗣️ {title}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.italic = True
    run_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def create_full_scope_10_slide_docx(output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    styles = doc.styles
    normal_font = styles['Normal'].font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("PROJECT REVIEW – I (FULL MAJOR PROJECT SCOPE)")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run("10-Slide Complete Proposal Slide Deck Script & Detailed Speaker Notes\nLikhibi: NLP Resources, Word Prediction, Machine Translation & Mobile Keyboard")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Domain:", "Natural Language Processing (NLP) / Machine Translation / Mobile Systems"),
        ("Project Coordinators:", "Mr. Nzanthung Odyuo | Mr. Nokshangthemba")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_k, cell_v = row.cells[0], row.cells[1]
        cell_k.width, cell_v.width = Inches(2.0), Inches(4.5)
        set_cell_background(cell_k, "F7FAFC")
        set_cell_background(cell_v, "FFFFFF")
        
        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(2)
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(9.5)
        rk.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(2)
        rv = pv.add_run(v)
        rv.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    slides = [
        {
            "num": "SLIDE 1",
            "title": "Title Slide (Full Scope Project Metadata)",
            "bullets": [
                "Title: Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole",
                "Review: Full Major Project Proposal Defense (Project Review – I)",
                "Domain: Natural Language Processing / Machine Translation / Computational Linguistics",
                "Target Language: Nagamese Creole (Lingua Franca of Nagaland)",
                "Presenter: [Your Name] | Roll No / USN: [Your Roll Number] | Dept of CSE",
                "Project Coordinators: Mr. Nzanthung Odyuo & Mr. Nokshangthemba"
            ],
            "notes": "Respected Project Coordinators Mr. Nzanthung Odyuo sir, Mr. Nokshangthemba sir, and evaluation committee members. I am [Your Name], presenting my full major project proposal: 'Likhibi: Computational Resource Curation, Contextual Language Modeling, and Prototype Neural Translation for Nagamese Creole'. Nagamese is the primary lingua franca across Nagaland, but remains classified as a Low-Resource Language. This project builds the complete NLP resource stack for Nagamese—including a 20,000-entry dictionary, parallel corpus, N-gram prediction engine, baseline Neural Machine Translation, and an offline Android keyboard interface."
        },
        {
            "num": "SLIDE 2",
            "title": "Executive Table of Contents",
            "bullets": [
                "1. Project Core (Aim, Objectives & Motivation)",
                "2. Literature Survey & Gap Analysis",
                "3. Problem Statement",
                "4. Overall Proposed Methodology & Two-Tier System Architecture",
                "5. Data Acquisition, Preprocessing & 20,000 Lexical Database",
                "6. Core NLP Component 1: Contextual Word Prediction Engine (N-Gram + Trie)",
                "7. Core NLP Component 2 & 3: Neural Machine Translation (NMT) & Android IME Platform",
                "8. Step-by-Step Project Workflow & Implementation Timeline",
                "9. Expected Project Outcomes & Quantitative Evaluation Metrics",
                "10. Key References"
            ],
            "notes": "Here is the 10-slide roadmap for the complete project proposal. I will establish our core aim and motivation, review literature gaps, define the problem statement, and present our full methodology covering dataset creation, word prediction, neural machine translation, Android deployment, and evaluation metrics."
        },
        {
            "num": "SLIDE 3",
            "title": "Aim, Objectives & Motivation (Core Summary)",
            "bullets": [
                "AIM: Create standardized NLP resources for Nagamese Creole, train an offline contextual word-prediction engine, construct an English <-> Nagamese Neural Machine Translation prototype, and deploy them on an Android Input Method Editor (IME).",
                "CORE OBJECTIVES (ENTIRE LIFECYCLE):",
                "  • Dataset Creation: Build 6,965-line monolingual corpus & 6,965-pair parallel corpus.",
                "  • Lexical Database: Compile a verified 20,000+ entry Nagamese digital dictionary (nagamese_lexicon.json).",
                "  • Word Prediction Engine: Train statistical N-gram language models + Trie prefix tree index.",
                "  • Neural Machine Translation: Implement baseline seq2seq / Transformer NMT model for Nagamese <-> English translation.",
                "  • Android Platform: Integrate models into an offline Android keyboard app operating under 5 MB RAM.",
                "MOTIVATION:",
                "  • Spoken by 30M+ people, yet ZERO native digital NLP tools or predictive keyboards exist in Roman script.",
                "  • Bridges severe regional 'data poverty' and enables digital access for Northeastern languages."
            ],
            "notes": "Slide 3 presents our project foundations: Our Aim is to build the full computational NLP stack for Nagamese—ranging from datasets and dictionaries to text prediction, machine translation, and a mobile keyboard. Our Objectives cover 5 key milestones: corpus building, a 20,000-entry dictionary, N-gram prediction models, Neural Machine Translation, and Android IME integration. Our Motivation stems from sociolinguistic reality: Nagamese is spoken by millions daily, yet smartphones provide zero native support. We are building the foundational infrastructure to solve this."
        },
        {
            "num": "SLIDE 4",
            "title": "Literature Survey & Identified Research Gaps",
            "bullets": [
                "1. Nagamese Linguistic Foundations (Sreedhar 1974, Baishya 2013, Boruah 2018): Documented Nagamese grammar, noun/verb compounding. -> GAP: Purely descriptive academic studies with ZERO computational datasets, tokenizers, or code artifacts.",
                "2. Indian Low-Resource NLP & Machine Translation (Joshi et al. 2020, Kumar et al. 2021): Highlighted severe 'Data Poverty' in regional Indian languages. -> GAP: Northeastern creoles are completely missing from Indic NLP benchmarks (e.g., IndicGLUE, Samanantar).",
                "3. Mobile Predictive Input & Seq2Seq Modeling (Fowler 2015, Bahdanau 2015, Vaswani 2017): Established Trie prefix indexing for mobile input and Attention/Transformer architectures for NMT. -> GAP: Zero implementations exist for Nagamese <-> English."
            ],
            "notes": "Literature shows three critical gaps: First, Nagamese linguistics research is purely descriptive without digital datasets. Second, Indian NLP benchmarks completely exclude Northeastern creoles. Third, mobile predictive input and machine translation architectures have never been built for Nagamese. Our project directly solves all three gaps across its full scope."
        },
        {
            "num": "SLIDE 5",
            "title": "Problem Statement",
            "bullets": [
                "Problem Definition: Existing mobile operating systems and computational platforms completely lack native digital lexicons, language models, and machine translation systems for Nagamese Creole.",
                "Key Impacts:",
                "  1. Severe typing friction and aggressive English/Hindi auto-correct errors for Nagamese mobile users.",
                "  2. Complete digital exclusion of Nagamese from automated translation tools.",
                "  3. Total absence of standardized computational datasets needed for downstream NLP research."
            ],
            "notes": "The problem statement is three-fold: First, Nagamese mobile users face constant typing friction due to auto-correct errors. Second, Nagamese is digitally excluded from translation platforms like Google Translate. Third, the research community lacks standardized digital corpora to build language technologies."
        },
        {
            "num": "SLIDE 6",
            "title": "Proposed Methodology – System Architecture",
            "bullets": [
                "Two-Tier System Architecture Overview:",
                "TIER 1: NLP RESEARCH PIPELINE (Python)",
                "  • Text Preprocessor -> Lexical DB Builder (20k entries) -> Parallel Corpus Aligner (6.9k pairs)",
                "  • Model Construction: Trie Indexer | N-Gram LM Trainer | Seq2Seq NMT Model Prototype",
                "TIER 2: ANDROID DEMONSTRATION PLATFORM (Kotlin)",
                "  • Android IME Layer (LikhibiImeService, CustomKeyboardView, Settings)",
                "  • On-Device Abstractions (PredictionEngine, TranslationEngine, LexicalRepository)",
                "Decoupling Principle: Python pipeline handles research & training; Android IME acts strictly as mobile host."
            ],
            "notes": "Slide 6 shows our two-tier system architecture: Tier 1 is the Python NLP Research Pipeline—handling text preprocessing, 20k dictionary compilation, parallel corpus alignment, N-gram language modeling, Trie indexing, and Neural Machine Translation prototyping. Tier 2 is the Android Demonstration Platform—a native Kotlin keyboard app that hosts the exported model assets offline."
        },
        {
            "num": "SLIDE 7",
            "title": "Proposed Methodology – Data & Lexical DB",
            "bullets": [
                "Custom Data Collection:",
                "  • Monolingual Corpus: Extracted text from 26 Nagamese scripture publications (6,965 sentences, ~185k tokens).",
                "  • Scraped Data: Web glossaries and community lexicons (xobdo.org, nagamesekhobor.com).",
                "  • Parallel Translation Corpus: 6,965 aligned English-Nagamese sentence pairs (bible_parallel_corpus.tsv).",
                "  • Code-Switching Dataset: 1,000+ high-frequency English/Hindi loanwords (school, phone, office, time).",
                "Preprocessing & Morphological Pipeline:",
                "  • Custom regex tokenizer tailored for Romanized creole text with boundary markers (<s>, </s>).",
                "  • Creole morphology inflections: Noun cases (-khan, -laga, -ke, -pora, -te) & Verb aspects (-se, -bo, -bole, -ina).",
                "  • Anti-Synthetic Purge Filter: Explicit rules blocking invalid compounds (e.g., homolaga).",
                "20,000+ Entry Lexical Database (nagamese_lexicon.json):",
                "  • Schema: Lemma, IPA phonetics, POS category, English definition, frequency, etymology tag.",
                "  • Ratio: 95.2% Native Nagamese + 3.8% English Loanwords + 1.0% Hindi Borrowings. 100% automated validation scan."
            ],
            "notes": "On Slide 7, we detail the data engine: We created a 6,965-line monolingual corpus, a 6,965-pair parallel corpus, and a 1,000-word code-switching dataset. Preprocessing includes custom regex tokenization and creole morphological rules with an anti-synthetic filter. The output is a 20,000+ entry verified dictionary in JSON schema with etymology tagging."
        },
        {
            "num": "SLIDE 8",
            "title": "Component 1: Contextual Prediction Engine",
            "bullets": [
                "Statistical N-Gram Language Model:",
                "  • Trains Unigram, Bigram, and Trigram probabilities with Add-k Smoothing: P(w_i | w_{i-1}) = (C(w_{i-1}, w_i) + k) / (C(w_{i-1}) + k * |V|).",
                "  • Uses Backoff smoothing for unseen sequences (Trigram -> Bigram -> Unigram).",
                "  • Model Perplexity Metric (PP): Evaluates predictive certainty on test sequences (Target: PP < 50).",
                "Character-Level Trie Prefix Tree:",
                "  • Character tree indexing 20,000+ dictionary entries with O(L) search latency (<1 ms).",
                "  • Serialized to compact JSON payload (~0.77 MB).",
                "Hybrid Reranking Engine:",
                "  • User typing prefix ('ja') -> Trie matches candidates ['jabo', 'jai', 'jani'].",
                "  • Context exists ('moi') -> N-gram context probability boosts ranking: FinalScore(w) = TrieFreq(w) + alpha * P_Ngram(w | Context)."
            ],
            "notes": "Slide 8 details Component 1—Word Prediction: We train N-gram models with Add-k smoothing and backoff to predict next words given context. We build a Trie prefix tree for O(L) sub-millisecond prefix completion. A hybrid reranker combines Trie candidate frequencies with N-gram context probabilities to surface top predictions."
        },
        {
            "num": "SLIDE 9",
            "title": "Component 2 & 3: Translation & Android Platform",
            "bullets": [
                "Component 2 — Prototype Neural Machine Translation (NMT):",
                "  • Task: Bidirectional English <-> Nagamese sentence translation.",
                "  • Architecture: Sequence-to-Sequence (Seq2Seq) model with Attention / Lightweight Transformer.",
                "  • Dataset: 6,965 aligned sentence pairs (bible_parallel_corpus.tsv).",
                "  • Output: Quantized ONNX / TFLite weights for translation inferencing.",
                "Component 3 — Android Demonstration Platform (IME Keyboard):",
                "  • Base Service: Built on native Android InputMethodService API.",
                "  • On-Device Abstraction: PredictionEngine, TranslationEngine, and LexicalRepository.",
                "  • Constraints: 100% Offline operation, <5 MB memory footprint, <10 ms suggestion response time.",
                "Implementation Timeline (Phases I - V): Corpus Extraction -> Lexical DB -> Prediction Engine -> NMT Prototyping -> IME Integration."
            ],
            "notes": "Slide 9 covers Components 2 and 3: Component 2 is a baseline Neural Machine Translation model trained on our 6,965 sentence pairs using Seq2Seq with Attention, exported to ONNX/TFLite format. Component 3 is the Android Keyboard platform hosting both prediction and translation models offline under 5 MB RAM. The timeline spans 5 clear implementation phases."
        },
        {
            "num": "SLIDE 10",
            "title": "Expected Outcomes, Metrics & References",
            "bullets": [
                "Complete Project Deliverables:",
                "  • 20,000+ Entry Validated Dictionary (nagamese_lexicon.json).",
                "  • Parallel Translation Corpus (bible_parallel_corpus.tsv, 6,965 pairs).",
                "  • Serialized Prediction Models (trie_index.json, bigrams.json).",
                "  • Prototype Neural Machine Translation Model (Nagamese <-> English).",
                "  • Working Offline Android Keyboard Application (APK).",
                "Quantitative Evaluation Metrics:",
                "  • Perplexity (PP): Measures prediction uncertainty (Target: PP < 50).",
                "  • Keystroke Savings (KS %): Percentage of typing keypresses saved.",
                "  • Translation BLEU Score: Evaluates machine translation quality against ground truth.",
                "  • Hardware Performance: Latency (<10 ms) and Memory Footprint (<10 MB RAM).",
                "Key References:",
                "  • Sreedhar (1974) - Nagamese | Baishya (2013) - Morphology | Joshi et al. (2020) - Low-Resource NLP",
                "  • Fowler et al. (2015) - Trie IME | Bahdanau et al. (2015) - Seq2Seq NMT | Vaswani et al. (2017) - Transformers"
            ],
            "notes": "Slide 10 summarizes project deliverables and evaluation metrics: Deliverables include the 20k dictionary, parallel corpus, prediction models, translation prototype, and Android APK. Evaluation uses Perplexity, Keystroke Savings, BLEU score for translation quality, and mobile hardware benchmarks. Thank you, Nzanthung sir, Nokshangthemba sir, and committee members. I welcome your questions."
        }
    ]

    for slide in slides:
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(3)
        
        r_num = h2.add_run(f"{slide['num']}: ")
        r_num.bold = True
        r_num.font.size = Pt(13)
        r_num.font.color.rgb = RGBColor(0x31, 0x82, 0xCE)
        
        r_title = h2.add_run(slide['title'])
        r_title.bold = True
        r_title.font.size = Pt(13)
        r_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        for bullet in slide['bullets']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            brun = bp.add_run(bullet)
            brun.font.size = Pt(9.5)

        add_callout_box(doc, "Speaker Notes (What to say)", slide['notes'])

    doc.save(output_path)
    print(f"Successfully generated Full Scope 10-Slide DOCX at: {output_path}")

if __name__ == "__main__":
    out = "f:/likhibi-main/docs/Nagamese_NLP_Full_Project_Proposal.docx"
    create_full_scope_10_slide_docx(out)
